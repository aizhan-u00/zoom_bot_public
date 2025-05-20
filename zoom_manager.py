"""
Module for managing the Zoom API for scheduling and managing meetings.

Provides the `ZoomManager` class for interacting with the Zoom API, including creating 
meetings, checking availability, deleting, and downloading recordings. Supports 
multiple Zoom accounts and suggests alternative time slots when none are available.
"""
import re
import base64
import datetime
from typing import Dict, Optional, List, Any, Tuple
import requests
from docx import Document
from dateutil.parser import parse
import pytz
from logger import logger
from config import ZOOM_API_BASE, ZOOM_ACCOUNTS, TIMEZONE

class ZoomManager:
    """Manages interactions with the Zoom API for meeting operations."""

    def __init__(self):
        """Initializes ZoomManager with a token cache."""
        self.last_error: Optional[str] = None
        self.token_cache: Dict[str, Tuple[str, datetime.datetime]] = {}  # email: (token, expiration time)
        self.timezone = pytz.timezone(TIMEZONE)
        logger.info("ZoomManager initialized")

    def get_last_error(self) -> Optional[str]:
        """Returns the last error message.

        Returns:
            Optional[str]: Last error message or None.
        """
        return self.last_error

    def _clear_last_error(self) -> None:
        """Clears the last error message."""
        self.last_error = None

    def get_access_token(self, account: Dict[str, str]) -> Optional[str]:
        """Retrieves or refreshes the OAuth access token for a Zoom account.

        Args:
            account (Dict[str, str]): Dictionary with client_id, client_secret, account_id, email.

        Returns:
            Optional[str]: Access token if successful, otherwise None.
        """
        self._clear_last_error()
        email = account['email']

        # Checking cache
        if email in self.token_cache:
            token, expiry = self.token_cache[email]
            if expiry > datetime.datetime.now(datetime.timezone.utc) + datetime.timedelta(minutes=5):
                logger.debug("Using cached token for %s", email)
                return token

        logger.info("Retrieving new access token for %s", email)
        try:
            credentials = f"{account['client_id']}:{account['client_secret']}"
            encoded_credentials = base64.b64encode(credentials.encode()).decode()

            response = requests.post(
                "https://api.zoom.us/oauth/token",
                headers={
                    "Authorization": f"Basic {encoded_credentials}",
                    "Content-Type": "application/x-www-form-urlencoded",
                },
                data={"grant_type": "account_credentials", "account_id": account['account_id']},
                timeout=10
            )
            response.raise_for_status()

            token_data = response.json()
            access_token = token_data.get("access_token")
            expires_in = token_data.get("expires_in", 3600)

            if not access_token:
                self.last_error = f"Access token not received for {email}"
                logger.error(self.last_error)
                return None

            # Caching token
            expiry = datetime.datetime.now(datetime.timezone.utc) + datetime.timedelta(
                seconds=expires_in
            )
            self.token_cache[email] = (access_token, expiry)
            logger.info("Access token received and cached for %s", email)
            return access_token

        except requests.RequestException as error:
            self.last_error = f"Error retrieving access token for {email}: {str(error)}"
            logger.error(self.last_error)
            return None

    def check_availability(self, account: Dict[str, str], start_time: datetime.datetime,
                         duration: int) -> bool:
        """Checks the availability of an account for a time slot.

        Args:
            account (Dict[str, str]): Account data with token and email.
            start_time (datetime.datetime): Meeting start time (UTC).
            duration (int): Meeting duration in minutes.

        Returns:
            bool: True, if available, False, in case of time conflict.
        """
        self._clear_last_error()
        logger.info("Checking availability for %s at %s", account['email'], start_time)

        headers = {
            "Authorization": f"Bearer {account['access_token']}", 
            "Content-Type": "application/json"
        }
        url = f"{ZOOM_API_BASE}/users/{account['email']}/meetings"

        try:
            response = requests.get(url, headers=headers, timeout=10)
            response.raise_for_status()

            adjusted_start = start_time - datetime.timedelta(minutes=30)
            end_time = adjusted_start + datetime.timedelta(minutes=duration + 60)

            meetings = response.json().get("meetings", [])
            logger.info("Found %d scheduled meetings for %s", len(meetings), account['email'])

            for meeting in meetings:
                meeting_start = parse(meeting["start_time"]).replace(tzinfo=datetime.timezone.utc)
                meeting_duration = meeting.get("duration", 0)
                meeting_end = meeting_start + datetime.timedelta(minutes=meeting_duration)

                if meeting_end >= adjusted_start and meeting_start <= end_time:
                    logger.info("Detected time conflict for %s", account['email'])
                    return False

            logger.info("Account %s is available", account['email'])
            return True

        except requests.RequestException as error:
            self.last_error = f"Error checking availability for {account['email']}: {str(error)}"
            logger.error(self.last_error)
            return False

    def find_alternative_slots(self, requested_date: datetime.date, duration: int) -> List[str]:
        """Finds available time slots on the specified day.

        Args:
            requested_date (datetime.date): Date for availability check.
            duration (int): Meeting duration in minutes.

        Returns:
            List[str]: List of available slots in HH:MM format.
        """
        self._clear_last_error()
        logger.info("Searching for alternative slots for %s", requested_date)
        available_slots = []

        # Checking for slots between 9:00 and 22:00 with 30 minut interval
        start_hour = 9
        end_hour = 22
        for hour in range(start_hour, end_hour):
            for minute in (0, 30):
                slot_time = datetime.time(hour, minute)
                slot_datetime = datetime.datetime.combine(requested_date, slot_time)
                slot_datetime = self.timezone.localize(slot_datetime).astimezone(pytz.UTC)

                for account in ZOOM_ACCOUNTS:
                    account["access_token"] = self.get_access_token(account)
                    if not account["access_token"]:
                        continue
                    if self.check_availability(account, slot_datetime, duration):
                        available_slots.append(slot_time.strftime("%H:%M"))
                        break  # Slots are found, not checking other accounts

        if not available_slots:
            self.last_error = f"No available slots on {requested_date} for duration {duration} minutes"
            logger.warning(self.last_error)

        logger.info("Found %d alternative slots", len(available_slots))
        return available_slots

    def create_meeting(self, account: Dict[str, str], meeting_data: Dict[str, Any]) -> Optional[Dict]:
        """Creates a new Zoom meeting.

        Args:
            account (Dict[str, str]): Account data with token and email.
            meeting_data (Dict[str, Any]): Meeting configuration data.

        Returns:
            Optional[Dict]: Meeting details, if created, otherwise None.
        """
        self._clear_last_error()
        logger.info("Creating meeting for %s", account['email'])

        headers = {
            "Authorization": f"Bearer {account['access_token']}",
            "Content-Type": "application/json; charset=UTF-8"
        }
        url = f"{ZOOM_API_BASE}/users/{account['email']}/meetings"

        try:
            response = requests.post(url, json=meeting_data, headers=headers, timeout=10)
            if response.status_code == 201:
                logger.info("Meeting successfully created for %s", account['email'])
                return response.json()

            self.last_error = f"Error creating meeting for {account['email']}: {response.status_code} - {response.text}"
            logger.error(self.last_error)
            return None

        except requests.RequestException as error:
            self.last_error = f"Error creating meeting for {account['email']}: {str(error)}"
            logger.error(self.last_error)
            return None

    def book_meeting(self, meeting_data: Dict[str, Any], requested_date: datetime.date) -> Tuple[Optional[Dict], List[str]]:
        """Books a meeting or suggests alternative slots if none are available.

        Args:
            meeting_data (Dict[str, Any]): Meeting configuration data.
            requested_date (datetime.date): Requested date for alternative slots.

        Returns:
            Tuple[Optional[Dict], List[str]]: Meeting details, if created, and list of alternative slots.
        """
        self._clear_last_error()
        logger.info("Starting meeting booking process")

        try:
            start_time = datetime.datetime.strptime(
                meeting_data['start_time'], "%Y-%m-%dT%H:%M:%SZ"
            ).replace(tzinfo=datetime.timezone.utc)
            duration = meeting_data['duration']

            meeting_data.update({
                "type": 2,  # Planned meeting
                "timezone": TIMEZONE,
                "settings": {
                    "host_video": False,
                    "participant_video": True,
                    "waiting_room": False,
                    "auto_recording": "cloud",
                    "meeting_authentication": False,
                    "join_before_host": True,
                    "jbh_time": 5,
                    "auto_start_meeting_summary": True,
                    "mute_upon_entry": True
                }
            })

            error_messages = []
            for account in ZOOM_ACCOUNTS:
                logger.info("Attempting with account: %s", account['email'])
                account["access_token"] = self.get_access_token(account)

                if not account["access_token"]:
                    error_messages.append(
                        f"Error retrieving token for {account['email']}: {self.last_error}" 
                    )
                    continue

                if self.check_availability(account, start_time, duration):
                    result = self.create_meeting(account, meeting_data)
                    if result:
                        logger.info("Meeting successfully booked with %s", account['email'])
                        return result, []
                    error_messages.append(
                        f"Error creating meeting for {account['email']}: {self.last_error}"
                    )
                else:
                    error_messages.append(
                        f"No available slots for {account['email']}: time conflict"
                    )

            # No available slots, looking for alternatives
            alternative_slots = self.find_alternative_slots(requested_date, duration)
            self.last_error = "\n".join(error_messages) if error_messages else "No available accounts"
            logger.warning("Booking failed: %s", self.last_error)
            return None, alternative_slots

        except ValueError as error:
            self.last_error = f"Invalid meeting data: {str(error)}"
            logger.error(self.last_error)
            return None, []

    def extract_meeting_id(self, zoom_link: str) -> Optional[str]:
        """Extracts the meeting ID from a Zoom link.

        Args:
            zoom_link (str): URL of Zoom meeting.

        Returns:
            Optional[str]: Meeting ID, if found, otherwise None.
        """
        self._clear_last_error()
        logger.info("Extracting meeting ID from URL: %s", zoom_link)

        match = re.search(r"/j/(\d+)(?:\?pwd=[^ ]*)?", zoom_link)
        if match:
            meeting_id = match.group(1)
            logger.debug("Extracted meeting ID: %s", meeting_id)
            return meeting_id

        self.last_error = "Invalid Zoom URL format"
        logger.error(self.last_error)
        return None

    def delete_meeting(self, meeting_url: str) -> Optional[str]:
        """Deletes a Zoom meeting by its URL.

        Args:
            meeting_url (str): URL of Zoom meeting.

        Returns:
            Optional[str]: Email of an account from which meeting is deleted, otherwise None.
        """
        self._clear_last_error()
        logger.info("Starting meeting deletion: %s", meeting_url)

        meeting_id = self.extract_meeting_id(meeting_url)
        if not meeting_id:
            self.last_error = "Error retrieving meeting ID from URL"
            return None

        url = f"{ZOOM_API_BASE}/meetings/{meeting_id}"
        error_messages = []

        for account in ZOOM_ACCOUNTS:
            logger.info("Checking account: %s", account['email'])
            account["access_token"] = self.get_access_token(account)

            if not account["access_token"]:
                error_messages.append(
                    f"Error retrieving token for {account['email']}: {self.last_error}"
                )
                continue

            headers = {
                "Authorization": f"Bearer {account['access_token']}",
                "Content-Type": "application/json"
            }

            try:
                response = requests.delete(url, headers=headers, timeout=10)
                if response.status_code == 204:
                    logger.info("Meeting deleted from account: %s", account['email'])
                    return account['email']

            except requests.RequestException as error:
                error_messages.append(
                    f"Error deleting meeting from {account['email']}: {str(error)}"
                )

        self.last_error = "\n".join(error_messages) if error_messages else "Meeting not found in accounts"
        logger.warning(self.last_error)
        return None

    def get_recording_url(self, meeting_id: str, specific_account: Optional[str] = None) -> Tuple[Optional[Dict], Optional[str], Optional[str]]:
        """Retrieves the recording URL for a meeting.

        Args:
            meeting_id (str): ID of Zoom meeting.
            specific_account (Optional[str]): Email of account for checking, if specified.

        Returns:
            Tuple[Optional[Dict], Optional[str], Optional[str]]: (account, download URL, topic) or (None, None, None).
        """
        self._clear_last_error()
        logger.info("Retrieving recording URL for meeting ID: %s", meeting_id)

        accounts_to_check = []
        if specific_account:
            for account in ZOOM_ACCOUNTS:
                if account['email'] == specific_account:
                    accounts_to_check = [account]
                    break
            if not accounts_to_check:
                self.last_error = f"Specified account {specific_account} not found"
                logger.error(self.last_error)
                return None, None, None
        else:
            accounts_to_check = ZOOM_ACCOUNTS

        error_messages = []
        for account in accounts_to_check:
            account["access_token"] = self.get_access_token(account)
            if not account["access_token"]:
                error_messages.append(
                    f"Error retrieving token for {account['email']}: {self.last_error}"
                )
                continue

            headers = {
                "Authorization": f"Bearer {account['access_token']}",
                "Content-Type": "application/json"
            }
            url = f"{ZOOM_API_BASE}/meetings/{meeting_id}/recordings"

            try:
                response = requests.get(url, headers=headers, timeout=10)
                if response.status_code == 200:
                    data = response.json()
                    download_url = None

                    for rec in data.get("recording_files", []):
                        if rec["file_extension"] == "MP4":
                            download_url = rec["download_url"]
                            break

                    if download_url:
                        self.download_summary(data["uuid"], account['access_token'], data["topic"])
                        logger.info("Recording found for account: %s", account['email'])
                        return account, download_url, data["topic"]

                error_messages.append(
                    f"Meeting nor found for {account['email']}: Status {response.status_code}"
                )
            except requests.RequestException as error:
                error_messages.append(
                    f"Error retrieving recording for {account['email']}: {str(error)}"
                )

        self.last_error = "\n".join(error_messages) if error_messages else "Recording not found"
        logger.warning(self.last_error)
        return None, None, None

    def delete_recording(self, meeting_id: str, token: str) -> None:
        """Deletes the meeting recording and summary.

        Args:
            meeting_id (str): ID of Zoom meeting.
            token (str): Access token for authentication.
        """
        self._clear_last_error()
        logger.info("Deleting recording for meeting ID: %s", meeting_id)

        headers = {"Authorization": f"Bearer {token}"}
        recording_url = f"{ZOOM_API_BASE}/meetings/{meeting_id}/recordings"
        summary_url = f"{ZOOM_API_BASE}/meetings/{meeting_id}/meeting_summary"

        try:
            response = requests.delete(recording_url, headers=headers, timeout=10)
            if response.status_code == 204:
                logger.info("Recording deleted: %s", meeting_id)
            else:
                self.last_error = f"Error deleting recording: {response.status_code} - {response.text}"
                logger.error(self.last_error)

            summary_response = requests.delete(summary_url, headers=headers, timeout=10)
            if summary_response.status_code == 204:
                logger.info("Summary deleted: %s", meeting_id)
            else:
                self.last_error = (
                    (self.last_error or "") + f"\nError deleting summary: {summary_response.status_code} - {summary_response.text}"
                )
                logger.error("Error deleting summary: %s", summary_response.status_code)

        except requests.RequestException as error:
            f"Deletion error: {str(error)}"
            logger.error(self.last_error)

    def download_summary(self, meeting_id: str, token: str, title: str) -> Optional[str]:
        """Downloads the meeting summary as a Word document.

        Args:
            meeting_id (str): ID of Zoom meeting.
            token (str): Access token for authentication.
            title (str): Meeting title.

        Returns:
            Optional[str]: Path to saved document, if successful, otherwise None.
        """
        self._clear_last_error()
        logger.info("Downloading summary for meeting: %s", title)

        url = f"{ZOOM_API_BASE}/meetings/{meeting_id}/meeting_summary"
        headers = {
            "Authorization": f"Bearer {token}",
            "Content-Type": "application/json"
        }

        try:
            response = requests.get(url, headers=headers, timeout=10)
            response.raise_for_status()

            data = response.json()
            main_summary = data.get("summary_overview", "Summary not available")
            summaries = data.get("summary_details", [])

            document = Document()
            document.add_heading(f'Meeting Summary: {title}', level=1)
            document.add_paragraph(main_summary)

            for summary in summaries:
                document.add_heading(summary.get("label", "Chapter"), level=2)
                document.add_paragraph(str(summary.get("summary", "")))

            file_path = f"{title}_summary.docx"
            document.save(file_path)

            logger.info("Summary saved to: %s", file_path)
            return file_path

        except requests.RequestException as error:
            self.last_error = f"Error downloading summary: {str(error)}"
            logger.error(self.last_error)
            return None
        except OSError as error:
            self.last_error = f"Error processing summary document: {str(error)}"
            logger.error(self.last_error)
            return None

    def download_recording(self, meeting_url: str, specific_account: Optional[str] = None) -> Optional[str]:
        """Downloads the meeting recording.

        Args:
            meeting_url (str): URL of Zoom meeting.
            specific_account (Optional[str]): Email of account for checking, if specified.

        Returns:
            Optional[str]: Meeting title, if downloaded successfully, otherwise None.
        """
        self._clear_last_error()
        logger.info("Downloading recording from: %s", meeting_url)

        meeting_id = self.extract_meeting_id(meeting_url)
        if not meeting_id:
            self.last_error = "Error retrieving meeting ID from URL"
            return None

        account, download_url, title = self.get_recording_url(meeting_id, specific_account)
        if not all((account, download_url, title)):
            self.last_error = self.last_error or "Recording not available"
            return None

        video_path = f"{title}.mp4"
        headers = {"Authorization": f"Bearer {account['access_token']}"}

        try:
            response = requests.get(download_url, headers=headers, stream=True, timeout=30)
            response.raise_for_status()

            logger.info("Downloading to: %s", video_path)
            with open(video_path, "wb") as file:
                for chunk in response.iter_content(chunk_size=8192):
                    file.write(chunk)

            logger.info("Download completed")
            self.delete_recording(meeting_id, account['access_token'])

            return title

        except requests.RequestException as error:
            self.last_error = f"Download error for account {account['email']}: {str(error)}"
            logger.error(self.last_error)
            return None
